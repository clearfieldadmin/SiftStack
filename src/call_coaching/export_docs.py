"""export_docs.py - turn call-coaching markdown reports into designed Word
documents and wire each per-call report to its actual recording.

Design system (v2, 2026-07-07):
  - Code fences are unwrapped, so the CALL GRADE REPORT block renders as real
    content (tables, headings), never as a monospace wall.
  - All pipe tables become real Word tables: navy shaded header row, right
    aligned numbers, bold shaded TOTAL row, 9pt.
  - Per-call docs open with a color-coded score banner (band color: green
    Strong/Elite, amber Developing, orange Needs Work, red Retrain/FAIL) and a
    "Listen to this call" callout linking the streaming MP3.
  - The metadata bullet block becomes a two-column info table.
  - ALL-CAPS section labels (SCORES, TOP 3 STRENGTHS...) become styled
    sub-headings; "> " quotes render italic gray and indented.
  - Number spacing normalized (62.5 /100 -> 62.5/100).

Also patches every per-call report .md with a Recording line (streaming URL +
local MP3 path) pulled from call_log.json. Idempotent.

USAGE (from SiftStack root, venv python):
  python src/call_coaching/export_docs.py             # patch + export everything
  python src/call_coaching/export_docs.py --no-patch  # export only
Output: output/call_coaching/reports/docx/{same relative path}.docx
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "output" / "call_coaching"
REPORTS = OUT / "reports"
DOCX_DIR = REPORTS / "docx"
SKIP_DIRS = {"docx", "_archive_wrong_rubric"}

NAVY = "1F3864"
NAVY_RGB = RGBColor(0x1F, 0x38, 0x64)
SLATE_RGB = RGBColor(0x40, 0x46, 0x52)
GRAY_RGB = RGBColor(0x6B, 0x70, 0x78)
LIGHT_FILL = "EEF3FA"
TOTAL_FILL = "DCE6F4"
BAND_COLORS = {
    "ELITE": "1E7B34", "STRONG": "2E7D32", "DEVELOPING": "B26A00",
    "NEEDS WORK": "C05000", "RETRAIN": "B71C1C", "FAIL": "B71C1C",
}


# ---------- data ----------

def call_map() -> dict:
    log = json.loads((OUT / "call_log.json").read_text(encoding="utf-8"))
    return {str(c["call_id"]): c for c in log if c.get("call_id")}


def patch_md_recording_links(calls: dict) -> int:
    patched = 0
    for md in REPORTS.rglob("*.md"):
        if any(part in SKIP_DIRS for part in md.relative_to(REPORTS).parts):
            continue
        m = re.match(r"^(\d{9,})_", md.name)
        if not m:
            continue
        call = calls.get(m.group(1))
        if not call or not call.get("recording_url"):
            continue
        text = md.read_text(encoding="utf-8")
        if "rec.smrtphone.io" in text:
            continue
        local = OUT / "recordings" / f"{call['call_id']}.mp3"
        line = (f"- Recording: [listen to this call]({call['recording_url']}) "
                f"(local file: {local})")
        lines = text.splitlines()
        insert_at = 1
        for i, ln in enumerate(lines):
            if ln.startswith("- "):
                insert_at = i + 1
            elif insert_at > 1 and not ln.startswith("- "):
                break
        lines.insert(insert_at, line)
        md.write_text("\n".join(lines) + "\n", encoding="utf-8")
        patched += 1
    return patched


# ---------- low-level docx helpers ----------

def shade_cell(cell, hex_fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_hyperlink(paragraph, url: str, text: str, size: Pt | None = None):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rpr.append(color)
    rpr.append(u)
    if size is not None:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size.pt * 2)))
        rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)\s]+\))")


def render_inline(paragraph, text: str, size: Pt | None = None, bold_all=False,
                  color: RGBColor | None = None, italic_all=False):
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
        elif chunk.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)\s]+)\)", chunk)
            if m:
                add_hyperlink(paragraph, m.group(2), m.group(1), size=size)
                continue
            run = paragraph.add_run(chunk)
        else:
            run = paragraph.add_run(chunk)
        if size is not None:
            run.font.size = size
        if bold_all:
            run.bold = True
        if italic_all:
            run.italic = True
        if color is not None:
            run.font.color.rgb = color


NUM_CELL = re.compile(r"^-?\$?\d[\d,.]*\s*(%|/100|/5|s|pts?|min)?$|^N/A.*$|^--$", re.I)


def normalize_numbers(text: str) -> str:
    text = re.sub(r"(\d(?:\.\d+)?)\s*/\s*100\b", r"\1/100", text)
    text = re.sub(r"(\d(?:\.\d+)?)\s*/\s*5\b", r"\1/5", text)
    text = re.sub(r"(\d)\s+%", r"\1%", text)
    return text


def band_of(text: str) -> str | None:
    for key in BAND_COLORS:
        if key in text.upper():
            return key
    return None


# ---------- document construction ----------

def style_base(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    for name, size in (("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11.5), ("Heading 4", 11)):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = NAVY_RGB
        st.paragraph_format.space_before = Pt(12 if name != "Heading 1" else 4)
        st.paragraph_format.space_after = Pt(4)
    for sec in doc.sections:
        sec.left_margin = sec.right_margin = Inches(0.8)
        sec.top_margin = sec.bottom_margin = Inches(0.7)


def add_table(doc, rows: list[list[str]]):
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        is_total = any("TOTAL" in (c or "").upper() for c in row[:2])
        for ci in range(ncols):
            cell = tbl.cell(ri, ci)
            txt = normalize_numbers(row[ci].strip()) if ci < len(row) else ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.space_before = Pt(1)
            render_inline(para, txt, size=Pt(9),
                          bold_all=(ri == 0 or is_total),
                          color=RGBColor(0xFF, 0xFF, 0xFF) if ri == 0 else None)
            if ri == 0:
                shade_cell(cell, NAVY)
            elif is_total:
                shade_cell(cell, TOTAL_FILL)
            if ri > 0 and ci > 0 and NUM_CELL.match(txt or ""):
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            band = band_of(txt) if ci > 0 else None
            if ri > 0 and band and len(txt) <= 24:
                for r_ in para.runs:
                    r_.font.color.rgb = RGBColor.from_string(BAND_COLORS[band])
                    r_.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_info_table(doc, pairs: list[tuple[str, str]]):
    tbl = doc.add_table(rows=len(pairs), cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(1.5)
    tbl.columns[1].width = Inches(5.4)
    for ri, (k, v) in enumerate(pairs):
        c0, c1 = tbl.cell(ri, 0), tbl.cell(ri, 1)
        c0.width, c1.width = Inches(1.5), Inches(5.4)
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(1)
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(9)
        r0.font.color.rgb = SLATE_RGB
        shade_cell(c0, LIGHT_FILL)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        render_inline(p1, normalize_numbers(v), size=Pt(9))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_score_banner(doc, score: str | None, band: str | None, autofail: str | None):
    if not (score or band):
        return
    fill = BAND_COLORS.get(band_of(band or "") or "", NAVY)
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    parts = []
    if score:
        parts.append(f"Score: {score}/100")
    if band:
        parts.append(f"Band: {band}")
    if autofail:
        parts.append(f"Auto-fail check: {autofail}")
    run = p.add_run("   |   ".join(parts))
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_listen_callout(doc, call: dict):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    shade_cell(cell, LIGHT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("▶  ")
    run.font.color.rgb = NAVY_RGB
    run.bold = True
    add_hyperlink(p, call["recording_url"], "Listen to this call (streaming MP3)", size=Pt(10.5))
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(f"Local file: {OUT / 'recordings' / (str(call['call_id']) + '.mp3')}")
    r2.font.size = Pt(8)
    r2.font.color.rgb = GRAY_RGB
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


CAPS_LABEL = re.compile(r"^([A-Z][A-Z0-9 ,#&/'-]{5,})(\s*\(.*\))?\s*:?\s*$")


def md_to_docx(md_path: Path, dest: Path, call: dict | None):
    raw = md_path.read_text(encoding="utf-8")
    # unwrap code fences: their content is structured report text, not code
    lines = [ln for ln in raw.splitlines() if not ln.strip().startswith("```")]

    full = "\n".join(lines)
    score_m = re.search(r"TOTAL[^\n]*?(\d+(?:\.\d+)?)\s*/\s*100", full)
    band_m = re.search(r"Grade band:\s*\[?([A-Za-z .#;()\d/-]+?)\]?\s*$", full, re.M)
    autofail_m = re.search(r"Auto-fail check:\s*(\S+)", full)

    doc = docx.Document()
    style_base(doc)

    table_buf: list[str] = []
    meta_buf: list[tuple[str, str]] = []
    banner_done = False
    title_done = False
    grade_line_re = re.compile(r"^(Grade band|Auto-fail check|Outcome|Call length|Caller):", re.I)

    def flush_table():
        nonlocal table_buf
        rows = [r for r in table_buf if not re.match(r"^\s*\|[\s:|-]+\|?\s*$", r)]
        if rows:
            add_table(doc, [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows])
        table_buf = []

    def flush_meta():
        nonlocal meta_buf, banner_done
        if meta_buf:
            add_info_table(doc, meta_buf)
            meta_buf = []
        if not banner_done:
            add_score_banner(doc,
                             score_m.group(1) if score_m else None,
                             band_m.group(1).strip() if band_m else None,
                             autofail_m.group(1) if autofail_m else None)
            banner_done = True

    for ln in lines:
        stripped = ln.strip()
        if stripped.startswith("|"):
            table_buf.append(stripped)
            continue
        if table_buf:
            flush_table()

        # metadata bullets right under the title -> info table
        if title_done and not banner_done:
            m = re.match(r"^-\s+\**([^:*]+)\**:?\**\s*:?\s*(.*)$", stripped) if stripped.startswith("- ") else None
            if m:
                key = m.group(1).strip().rstrip(":")
                if key.islower():
                    key = key[0].upper() + key[1:]
                if key.lower() == "recording" and call and call.get("recording_url"):
                    continue  # the listen callout above already carries this
                meta_buf.append((key, m.group(2).strip()))
                continue
            if stripped and meta_buf:
                flush_meta()

        if not stripped:
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            if meta_buf:
                flush_meta()
            level = min(len(m.group(1)), 4)
            doc.add_heading(normalize_numbers(re.sub(r"\*\*", "", m.group(2))), level=level)
            if level == 1 and not title_done:
                title_done = True
                if call and call.get("recording_url"):
                    add_listen_callout(doc, call)
            continue

        # ALL-CAPS section labels from the unwrapped grade block
        cm = CAPS_LABEL.match(stripped)
        if cm and not stripped.startswith("- "):
            if meta_buf:
                flush_meta()
            h = doc.add_heading(cm.group(1).title() + (cm.group(2) or ""), level=3)
            for r_ in h.runs:
                r_.font.color.rgb = SLATE_RGB
            continue

        # grade-block key lines get color treatment
        if grade_line_re.match(stripped):
            p = doc.add_paragraph()
            key, _, val = stripped.partition(":")
            kr = p.add_run(key + ": ")
            kr.bold = True
            val = normalize_numbers(val.strip())
            band = band_of(val)
            vr = p.add_run(val)
            if band:
                vr.font.color.rgb = RGBColor.from_string(BAND_COLORS[band])
                vr.bold = True
            elif "PASS" in val:
                vr.font.color.rgb = RGBColor.from_string(BAND_COLORS["STRONG"])
                vr.bold = True
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            render_inline(p, normalize_numbers(stripped[2:]), italic_all=True, color=SLATE_RGB)
            continue

        m = re.match(r"^\s*[-*]\s+(.*)", ln)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            render_inline(p, normalize_numbers(m.group(1)))
            continue
        m = re.match(r"^\s*(\d+)[.)]\s+(.*)", ln)
        if m:
            # literal numbers from the source, so each section's list starts at 1
            # (Word's List Number style would continue counting across sections)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            num = p.add_run(f"{m.group(1)}.  ")
            num.bold = True
            render_inline(p, normalize_numbers(m.group(2)))
            continue

        p = doc.add_paragraph()
        render_inline(p, normalize_numbers(stripped))

    if table_buf:
        flush_table()
    if meta_buf:
        flush_meta()
    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))


def main() -> int:
    ap = argparse.ArgumentParser(description="Export coaching reports to designed Word docs")
    ap.add_argument("--no-patch", action="store_true")
    ap.add_argument("--only", help="only files whose name contains this substring")
    args = ap.parse_args()

    calls = call_map()
    if not args.no_patch:
        n = patch_md_recording_links(calls)
        print(f"Patched recording links into {n} report .md files")

    exported = 0
    for md in sorted(REPORTS.rglob("*.md")):
        rel = md.relative_to(REPORTS)
        if any(part in SKIP_DIRS for part in rel.parts):
            continue
        if args.only and args.only not in md.name:
            continue
        m = re.match(r"^(\d{9,})_", md.name)
        call = calls.get(m.group(1)) if m else None
        md_to_docx(md, DOCX_DIR / rel.with_suffix(".docx"), call)
        exported += 1
    print(f"Exported {exported} .docx files -> {DOCX_DIR}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
